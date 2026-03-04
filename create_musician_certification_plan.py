from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# 创建 Word 文档
doc = Document()

# 设置标题样式
title_style = doc.styles['Title']
title_font = title_style.font
title_font.size = Pt(24)
title_font.bold = True

# 添加标题
title = doc.add_heading('撕歌 APP 官方音乐人认证活动策划方案', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加空行
doc.add_paragraph()

# 一、活动概述
doc.add_heading('一、活动概述', 1)

# 活动名称
doc.add_heading('活动名称', 2)
doc.add_paragraph('撕歌 APP 官方音乐人认证专场')

# 活动目的
doc.add_heading('活动目的', 2)
purposes = [
    '为优质音乐人提供官方认证机会',
    '提升平台音乐内容质量',
    '增强用户参与感和归属感',
    '发掘和培养平台音乐人才'
]
for purpose in purposes:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(purpose)

# 活动时间
doc.add_heading('活动时间', 2)
doc.add_paragraph('• 报名时间：活动发布后 3 天内')
doc.add_paragraph('• 审核时间：报名截止后 1-2 个工作日内（具体时间待定）')
doc.add_paragraph('• 结果公布：审核结束后 1-3 个工作日内')

# 活动形式
doc.add_heading('活动形式', 2)
doc.add_paragraph('官方主持人在专属房间进行实时审核，参赛歌手依次进入房间演唱，通过者获得官方音乐人认证标识。')

# 二、活动规则
doc.add_heading('二、活动规则', 1)

# 参赛要求
doc.add_heading('参赛要求', 2)
doc.add_paragraph('1. 参赛资格：撕歌 APP 注册用户，账号状态正常', style='List Number')
doc.add_paragraph('2. 演唱要求：', style='List Number')
doc.add_paragraph('   • 演唱时长控制在 1-2 分钟内', style='List Bullet')
doc.add_paragraph('   • 可选择演唱歌曲高潮部分或半首歌', style='List Bullet')
doc.add_paragraph('   • 必须清唱或使用伴奏（禁止播放原唱）', style='List Bullet')
doc.add_paragraph('   • 音质清晰，无明显杂音', style='List Bullet')
doc.add_paragraph('3. 内容要求：', style='List Number')
doc.add_paragraph('   • 歌曲内容健康向上，符合平台规范', style='List Bullet')
doc.add_paragraph('   • 不得包含敏感、违规或侵权内容', style='List Bullet')
doc.add_paragraph('   • 建议选择能展现个人特色的歌曲', style='List Bullet')

# 审核标准
doc.add_heading('审核标准', 2)
standards_table = doc.add_table(rows=5, cols=3)
standards_table.style = 'Table Grid'
standards_table.rows[0].cells[0].text = '评分项'
standards_table.rows[0].cells[1].text = '权重'
standards_table.rows[0].cells[2].text = '说明'

standards_data = [
    ('音准节奏', '40%', '音准准确，节奏稳定'),
    ('音色表现', '30%', '音色独特，有辨识度'),
    ('情感表达', '20%', '情感真挚，感染力强'),
    ('整体印象', '10%', '台风自然，表现力佳')
]

for i, (item, weight, desc) in enumerate(standards_data, 1):
    standards_table.rows[i].cells[0].text = item
    standards_table.rows[i].cells[1].text = weight
    standards_table.rows[i].cells[2].text = desc

# 三、活动流程
doc.add_heading('三、活动流程', 1)

# 第一阶段
doc.add_heading('第一阶段：报名阶段', 2)
doc.add_paragraph('1. 报名渠道：加入"撕歌官方音乐认证群"', style='List Number')
doc.add_paragraph('2. 报名方式：群内接龙报名', style='List Number')
doc.add_paragraph('3. 报名信息：', style='List Number')
doc.add_paragraph('   • 用户 ID', style='List Bullet')
doc.add_paragraph('   • 用户昵称', style='List Bullet')
doc.add_paragraph('   • 参赛歌曲名称', style='List Bullet')
doc.add_paragraph('   • 歌曲原唱者', style='List Bullet')
doc.add_paragraph('4. 报名限制：', style='List Number')
doc.add_paragraph('   • 限前 50 名报名用户（可根据实际情况调整）', style='List Bullet')
doc.add_paragraph('   • 每人限报 1 次', style='List Bullet')
doc.add_paragraph('   • 报名截止后不再接受报名', style='List Bullet')

# 第二阶段
doc.add_heading('第二阶段：审核阶段', 2)
doc.add_paragraph('1. 房间准备：', style='List Number')
doc.add_paragraph('   • 官方创建专属审核房间', style='List Bullet')
doc.add_paragraph('   • 房间名称："撕歌官方音乐人认证专场"', style='List Bullet')
doc.add_paragraph('   • 房间密码：仅对报名成功用户公布', style='List Bullet')
doc.add_paragraph('2. 审核流程：', style='List Number')
doc.add_paragraph('   • 主持人开场介绍活动规则（5 分钟）', style='List Bullet')
doc.add_paragraph('   • 按报名顺序依次邀请参赛者进入房间', style='List Bullet')
doc.add_paragraph('   • 每位参赛者演唱 1-2 分钟', style='List Bullet')
doc.add_paragraph('   • 主持人现场点评并记录结果', style='List Bullet')
doc.add_paragraph('   • 审核全程录音录像存档', style='List Bullet')
doc.add_paragraph('3. 时间控制：', style='List Number')
doc.add_paragraph('   • 总审核时长控制在 2-3 小时内', style='List Bullet')
doc.add_paragraph('   • 如遇技术问题可适当延长', style='List Bullet')

# 第三阶段
doc.add_heading('第三阶段：结果公布', 2)
doc.add_paragraph('1. 审核周期：1-3 个工作日', style='List Number')
doc.add_paragraph('2. 公布渠道：撕歌官方音乐认证群', style='List Number')
doc.add_paragraph('3. 公布内容：', style='List Number')
doc.add_paragraph('   • 通过认证的用户名单', style='List Bullet')
doc.add_paragraph('   • 获得的官方音乐人认证标识', style='List Bullet')
doc.add_paragraph('   • 后续权益说明', style='List Bullet')
doc.add_paragraph('4. 未通过处理：', style='List Number')
doc.add_paragraph('   • 可在下次活动重新报名', style='List Bullet')
doc.add_paragraph('   • 提供简要改进建议（如需要）', style='List Bullet')

# 四、参与渠道详情
doc.add_heading('四、参与渠道详情', 1)

doc.add_heading('官方音乐认证群', 2)
doc.add_paragraph('• 群名称：撕歌官方音乐认证群')
doc.add_paragraph('• 入群方式：扫描活动海报二维码或通过官方客服邀请')
doc.add_paragraph('• 群管理：')
doc.add_paragraph('   • 群主：官方运营人员')
doc.add_paragraph('   • 管理员：活动负责人、技术支持')
doc.add_paragraph('• 群规要求：')
doc.add_paragraph('   • 禁止广告、刷屏')
doc.add_paragraph('   • 文明交流，互相尊重')
doc.add_paragraph('   • 关注群公告，及时获取活动信息')

# 报名接龙格式
doc.add_heading('报名接龙格式', 2)
format_text = """【撕歌音乐人认证报名】
1. 用户 ID：123456789
   用户昵称：张三
   参赛歌曲：《平凡之路》
   原唱：朴树
   
2. 用户 ID：987654321
   用户昵称：李四
   参赛歌曲：《演员》
   原唱：薛之谦"""
doc.add_paragraph(format_text)

# 五、活动宣传
doc.add_heading('五、活动宣传', 1)

doc.add_heading('宣传渠道', 2)
channels = [
    'APP 内推送：首页 banner、消息中心',
    '社交媒体：官方微博、微信公众号',
    '社群传播：各用户群、音乐爱好者群',
    'KOL 推广：邀请已认证音乐人转发'
]
for channel in channels:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(channel)

doc.add_heading('宣传物料', 2)
materials = [
    '活动海报：包含活动时间、参与方式、奖励说明',
    '宣传视频：往期优秀认证音乐人展示',
    'FAQ 文档：常见问题解答'
]
for material in materials:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(material)

# 六、后续权益
doc.add_heading('六、后续权益', 1)

doc.add_heading('官方音乐人认证标识', 2)
benefits1 = [
    '个人主页显示专属认证徽章',
    '作品推荐权重提升',
    '优先参与官方活动',
    '专属客服通道'
]
for benefit in benefits1:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(benefit)

doc.add_heading('发展机会', 2)
benefits2 = [
    '推荐至合作音乐平台',
    '优先获得商业合作机会',
    '参与官方音乐制作项目',
    '定期技能培训和指导'
]
for benefit in benefits2:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(benefit)

# 七、风险预案
doc.add_heading('七、风险预案', 1)

doc.add_heading('技术风险', 2)
tech_risks = [
    '网络问题：准备备用房间，允许重试一次',
    '设备故障：提前测试设备，准备技术支持',
    '系统崩溃：延后审核时间，及时通知用户'
]
for risk in tech_risks:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(risk)

doc.add_heading('运营风险', 2)
ops_risks = [
    '报名超员：严格按照先到先得原则',
    '争议处理：建立申诉机制，保证公平公正',
    '负面舆情：及时回应，积极沟通解决'
]
for risk in ops_risks:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(risk)

# 八、预算规划
doc.add_heading('八、预算规划', 1)

doc.add_heading('人力成本', 2)
labor_costs = [
    '主持人：1-2 名官方人员',
    '技术支持：1 名技术人员',
    '客服人员：1-2 名'
]
for cost in labor_costs:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(cost)

doc.add_heading('物料成本', 2)
material_costs = [
    '宣传设计费用',
    '群管理工具费用',
    '认证标识制作费用'
]
for cost in material_costs:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(cost)

doc.add_heading('时间成本', 2)
time_costs = [
    '准备阶段：3-5 天',
    '执行阶段：1-2 天',
    '后续跟进：3-5 天'
]
for cost in time_costs:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(cost)

# 九、效果评估
doc.add_heading('九、效果评估', 1)

doc.add_heading('评估指标', 2)
metrics = [
    '报名人数及完成率',
    '用户满意度调查',
    '认证通过率',
    '后续活跃度提升',
    '新增用户增长'
]
for metric in metrics:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(metric)

doc.add_heading('优化方向', 2)
optimizations = [
    '根据反馈调整审核标准',
    '优化报名流程',
    '增加活动频次',
    '丰富认证权益'
]
for opt in optimizations:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(opt)

# 备注
doc.add_paragraph()
doc.add_paragraph('备注：本策划方案可根据实际运营情况进行调整，确保活动顺利进行并达到预期效果。')

# 保存文档
output_path = '/home/admin/.openclaw/workspace/撕歌 APP 官方音乐人认证活动策划.docx'
doc.save(output_path)
print(f"文档已保存至：{output_path}")
