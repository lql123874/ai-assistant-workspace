#!/usr/bin/env python3
"""
快速文档生成工具
支持 Markdown 转 Word，保持格式和表格
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
import sys

def markdown_to_word(md_content, output_path):
    """将 Markdown 内容转换为 Word 文档"""
    doc = Document()
    
    lines = md_content.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # 标题
        if line.startswith('# '):
            p = doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            p = doc.add_heading(line[4:], level=3)
        
        # 表格
        elif line.startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            i -= 1
            
            if len(table_lines) > 1:
                # 解析表格
                rows = []
                for tl in table_lines:
                    if '---' not in tl:  # 跳过分隔行
                        cells = [c.strip() for c in tl.split('|')[1:-1]]
                        rows.append(cells)
                
                if rows:
                    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
                    table.style = 'Light Grid Accent 1'
                    
                    for row_idx, row_data in enumerate(rows):
                        for col_idx, cell_data in enumerate(row_data):
                            table.rows[row_idx].cells[col_idx].text = cell_data
        
        # 列表
        elif line.startswith('- ') or line.startswith('* '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif re.match(r'^\d+\. ', line):
            doc.add_paragraph(re.sub(r'^\d+\. ', '', line), style='List Number')
        
        # 普通段落
        elif line.strip():
            doc.add_paragraph(line)
        
        i += 1
    
    doc.save(output_path)
    return output_path

if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: doc_gen.py <input.md> <output.docx>")
        sys.exit(1)
    
    with open(sys.argv[1], 'r', encoding='utf-8') as f:
        content = f.read()
    
    output = markdown_to_word(content, sys.argv[2])
    print(f"Generated: {output}")
