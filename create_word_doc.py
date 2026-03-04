from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

# 创建Word文档
doc = Document()

# 设置标题样式
title_style = doc.styles['Title']
title_font = title_style.font
title_font.size = Pt(24)
title_font.bold = True

# 添加标题
title = doc.add_heading('撕歌APP官方音乐人认证活动详细执行流程（Work文档）', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 添加副标题
subtitle = doc.add_paragraph('—— 完整版执行手册')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(14)
subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)

# 添加空行
doc.add_paragraph()

# 第一部分：前期准备阶段
doc.add_heading('一、前期准备阶段（活动前5-7天）', 1)

# 1.1 团队组建与分工
doc.add_heading('1.1 团队组建与分工', 2)
table = doc.add_table(rows=1, cols=4)
table.style = 'Table Grid'
hdr_cells = table.rows[0].cells
hdr_cells[0].text = '角色'
hdr_cells[1].text = '人员'
hdr_cells[2].text = '职责'
hdr_cells[3].text = '具体任务'

# 添加表格行
roles = [
    ('项目负责人', '1人', '整体统筹', '制定时间表、协调资源、监督进度'),
    ('主持人', '1-2人', '审核执行', '熟悉审核标准、准备话术、现场把控'),
    ('技术支持', '1人', '技术保障', '房间创建、设备测试、故障处理'),
    ('客服人员', '2人', '用户服务', '群管理、答疑解惑、信息收集'),
    ('宣传专员', '1人', '推广宣传', '海报设计、内容发布、数据跟踪')
]

for role in roles:
    row_cells = table.add_row().cells
    row_cells[0].text = role[0]
    row_cells[1].text = role[1]
    row_cells[2].text = role[2]
    row_cells[3].text = role[3]

# 1.2 资源准备清单
doc.add_heading('1.2 资源准备清单', 2)
checklist = [
    '创建官方音乐认证群（QQ群/微信群）',
    '设计活动海报（包含二维码、时间、规则）',
    '准备FAQ文档（常见问题解答）',
    '制作审核评分表（Excel模板）',
    '准备备用审核房间（技术预案）',
    '配置官方认证标识（系统后台）',
    '准备录音录像设备（存档用）'
]

for item in checklist:
    p = doc.add_paragraph(style='List Bullet')
    p.add_run(item)

# 1.3 宣传预热
doc.add_heading('1.3 宣传预热（活动前3天开始）', 2)
doc.add_paragraph('Day -3:')
doc.add_paragraph('- APP首页Banner上线', style='List Bullet')
doc.add_paragraph('- 官方社交媒体发布预告', style='List Bullet')
doc.add_paragraph('- 向现有音乐人群发邀请', style='List Bullet')

doc.add_paragraph('Day -2:')
doc.add_paragraph('- 发布详细活动规则', style='List Bullet')
doc.add_paragraph('- 开放官方认证群入群通道', style='List Bullet')
doc.add_paragraph('- KOL转发推广', style='List Bullet')

doc.add_paragraph('Day -1:')
doc.add_paragraph('- 最后提醒报名即将开始', style='List Bullet')
doc.add_paragraph('- 群内发布报名接龙模板', style='List Bullet')
doc.add_paragraph('- 技术设备最终测试', style='List Bullet')

# 第二部分：报名阶段
doc.add_heading('二、报名阶段（活动当天00:00 - 报名截止）', 1)

# 2.1 报名启动流程
doc.add_heading('2.1 报名启动流程', 2)
doc.add_paragraph('00:00 报名开启')
doc.add_paragraph('1. 客服在官方认证群发布公告：', style='List Number')
doc.add_paragraph('【撕歌官方音乐人认证活动报名正式开启！】', style='Intense Quote')
doc.add_paragraph('📅 报名时间：即日起至[具体日期] [具体时间]', style='Normal')
doc.add_paragraph('👥 名额限制：前50名（先到先得）', style='Normal')
doc.add_paragraph('📝 报名格式：严格按照以下模板接龙', style='Normal')
doc.add_paragraph('【撕歌音乐人认证报名】', style='Normal')
doc.add_paragraph('1. 用户ID：[你的撕歌APP用户ID]', style='Normal')
doc.add_paragraph('   用户昵称：[你的撕歌APP昵称]', style='Normal')
doc.add_paragraph('   参赛歌曲：[完整歌曲名称]', style='Normal')
doc.add_paragraph('   原唱：[原唱歌手姓名]', style='Normal')

doc.add_paragraph('2. 客服人员实时监控接龙情况', style='List Number')
doc.add_paragraph('3. 每10人更新一次剩余名额提醒', style='List Number')

# 继续添加更多内容...
doc.add_heading('2.2 报名审核与确认', 2)
doc.add_paragraph('实时审核（报名期间）')
doc.add_paragraph('- 检查用户ID格式是否正确', style='List Bullet')
doc.add_paragraph('- 验证账号状态是否正常', style='List Bullet')
doc.add_paragraph('- 确认歌曲内容符合规范', style='List Bullet')
doc.add_paragraph('- 标记有问题的报名信息', style='List Bullet')

doc.add_paragraph('报名截止后2小时内')
doc.add_paragraph('1. 整理最终报名名单（Excel表格）', style='List Number')
doc.add_paragraph('2. 发送确认消息给所有报名成功用户', style='List Number')
doc.add_paragraph('3. 私信通知报名失败用户（超员或信息错误）', style='List Number')

# 第三部分：审核执行阶段
doc.add_heading('三、审核执行阶段（报名截止后1-2天）', 1)

# 3.1 审核前准备
doc.add_heading('3.1 审核前准备（审核当天）', 2)
doc.add_paragraph('提前2小时：')
items = [
    '主持人、技术支持到位',
    '创建主审核房间和备用房间', 
    '测试麦克风、网络、录制设备',
    '准备审核评分表（打印/电子版）',
    '确认报名名单和演唱顺序'
]
for item in items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph('提前30分钟：')
items2 = [
    '在官方认证群公布房间密码',
    '提醒参与者提前进入等候',
    '发布审核注意事项'
]
for item in items2:
    doc.add_paragraph(item, style='List Bullet')

# 3.2 审核执行流程
doc.add_heading('3.2 审核执行流程（详细时间表）', 2)
doc.add_paragraph('单个选手审核SOP：')
doc.add_paragraph('1. 邀请进入（30秒）', style='List Number')
doc.add_paragraph('   - "下一位，请[昵称]进入房间"', style='List Bullet 2')
doc.add_paragraph('   - 等待选手进入并确认身份', style='List Bullet 2')

doc.add_paragraph('2. 演唱环节（60-90秒）', style='List Number')
doc.add_paragraph('   - "请开始您的演唱"', style='List Bullet 2')
doc.add_paragraph('   - 计时器启动', style='List Bullet 2')
doc.add_paragraph('   - 90秒时提醒："还有30秒"', style='List Bullet 2')
doc.add_paragraph('   - 120秒时打断："时间到，感谢演唱"', style='List Bullet 2')

doc.add_paragraph('3. 点评记录（30秒）', style='List Number')
doc.add_paragraph('   - 简要口头点评（1-2句话）', style='List Bullet 2')
doc.add_paragraph('   - 在评分表上打分', style='List Bullet 2')
doc.add_paragraph('   - "感谢参与，请退出房间"', style='List Bullet 2')

# 第四部分：结果处理阶段
doc.add_heading('四、结果处理阶段（审核结束后1-3个工作日）', 1)
doc.add_paragraph('结果公布时间：审核结束后第3个工作日 18:00')
doc.add_paragraph('公布内容示例：')
doc.add_paragraph('【撕歌官方音乐人认证活动结果公布】', style='Intense Quote')
doc.add_paragraph('🎉 恭喜以下用户通过官方音乐人认证！', style='Normal')
doc.add_paragraph('📋 通过名单：', style='Normal')
doc.add_paragraph('1. 用户ID：123456789 | 昵称：张三', style='Normal')
doc.add_paragraph('2. 用户ID：987654321 | 昵称：李四', style='Normal')
doc.add_paragraph('...', style='Normal')

# 第五部分：后续运营
doc.add_heading('五、后续运营阶段', 1)
doc.add_paragraph('认证用户管理：')
doc.add_paragraph('- 建立认证音乐人群：专门的交流群', style='List Bullet')
doc.add_paragraph('- 定期活动：每月主题演唱、技能分享', style='List Bullet')
doc.add_paragraph('- 商业对接：推荐合适的商业合作机会', style='List Bullet')
doc.add_paragraph('- 内容扶持：优先推荐优质作品', style='List Bullet')

# 第六部分：风险控制
doc.add_heading('六、风险控制与应急预案', 1)
doc.add_paragraph('主要风险点：')
doc.add_paragraph('• 报名超员：严格限流，提前说明', style='List Bullet')
doc.add_paragraph('• 技术故障：提前测试，备用方案', style='List Bullet')
doc.add_paragraph('• 争议投诉：标准透明，全程录像', style='List Bullet')
doc.add_paragraph('• 负面舆情：主动沟通，及时回应', style='List Bullet')

# 第七部分：工具与模板
doc.add_heading('七、工具与模板', 1)
doc.add_paragraph('必用工具清单：')
tools = [
    '报名管理：腾讯文档/石墨文档（实时协作）',
    '评分记录：Excel评分模板',
    '通讯工具：企业微信/钉钉（内部沟通）',
    '录制设备：OBS Studio（房间录制）',
    '数据分析：Google Analytics/友盟（效果追踪）'
]
for tool in tools:
    doc.add_paragraph(tool, style='List Bullet')

doc.add_paragraph('\n标准化模板：')
templates = [
    '报名接龙模板',
    '审核评分表模板',
    '结果通知模板', 
    'FAQ问答模板',
    '应急联系人清单'
]
for template in templates:
    doc.add_paragraph(template, style='List Bullet')

# 文档结尾
doc.add_paragraph('\n' + '='*50)
doc.add_paragraph('文档维护说明：')
doc.add_paragraph('- 本流程文档为动态更新版本', style='List Bullet')
doc.add_paragraph('- 每次活动后根据实际情况进行优化', style='List Bullet')
doc.add_paragraph('- 建议保存每次活动的执行记录作为参考', style='List Bullet')
doc.add_paragraph('- 重要修改需经项目负责人确认', style='List Bullet')

doc.add_paragraph('\n最后更新：2026年3月1日')
doc.add_paragraph('版本号：v1.0')

# 保存文档
doc.save('/home/admin/.openclaw/workspace/撕歌APP官方音乐人认证活动详细执行流程.docx')
print("Word文档创建成功！")