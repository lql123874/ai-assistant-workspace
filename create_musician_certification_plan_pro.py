from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn

# 创建 Word 文档
doc = Document()

# 设置页面
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1.25)
section.right_margin = Inches(1.25)

# ==================== 封面页 ====================
# 添加空行使内容居中
for _ in range(3):
    doc.add_paragraph()

# 活动名称（封面）
title = doc.add_heading('撕歌 APP 官方音乐人认证活动\n策划方案', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = RGBColor(0, 51, 102)

doc.add_paragraph()

# 副标题
subtitle = doc.add_paragraph('—— 发掘音乐人才 · 打造专业平台 ——')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.runs[0].font.size = Pt(16)
subtitle.runs[0].font.color.rgb = RGBColor(100, 100, 100)
subtitle.runs[0].font.italic = True

# 空行
for _ in range(8):
    doc.add_paragraph()

# 底部信息
info_para = doc.add_paragraph()
info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

info_para.add_run('主办单位：撕歌 APP 运营中心\n')
info_para.add_run('策划部门：音乐人生态部\n')
info_para.add_run('版本号：V1.0\n')
info_para.add_run('编制日期：2026 年 3 月 2 日')

for run in info_para.runs:
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(80, 80, 80)

# 分页符
doc.add_page_break()

# ==================== 目录页 ====================
toc_title = doc.add_heading('目  录', 1)
toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

toc_items = [
    ('一、执行摘要', 1),
    ('二、活动背景与目的', 2),
    ('    2.1 活动背景', 3),
    ('    2.2 活动目的', 3),
    ('    2.3 活动目标', 3),
    ('三、活动主题与定位', 2),
    ('四、活动概况', 2),
    ('    4.1 基本信息', 3),
    ('    4.2 参与对象', 3),
    ('    4.3 活动形式', 3),
    ('五、目标受众分析', 2),
    ('六、活动流程详细规划', 2),
    ('    6.1 第一阶段：前期准备（活动前 5-7 天）', 3),
    ('    6.2 第二阶段：宣传预热（活动前 3 天）', 3),
    ('    6.3 第三阶段：报名阶段（活动当天）', 3),
    ('    6.4 第四阶段：审核阶段（报名截止后 1-2 天）', 3),
    ('    6.5 第五阶段：结果公布（审核结束后 1-3 天）', 3),
    ('七、组织架构与人员分工', 2),
    ('    7.1 组织架构图', 3),
    ('    7.2 人员职责明细', 3),
    ('八、宣传推广方案', 2),
    ('    8.1 宣传渠道', 3),
    ('    8.2 宣传物料', 3),
    ('    8.3 宣传节奏', 3),
    ('九、预算明细表', 2),
    ('十、风险评估与应急预案', 2),
    ('    10.1 技术风险', 3),
    ('    10.2 运营风险', 3),
    ('    10.3 舆情风险', 3),
    ('十一、效果评估与 KPI', 2),
    ('十二、后续权益与发展规划', 2),
    ('附录', 1),
]

for i, (item, level) in enumerate(toc_items, 1):
    p = doc.add_paragraph()
    if level == 1:
        run = p.add_run(f'{i}. {item}')
        run.font.bold = True
        run.font.size = Pt(14)
    elif level == 2:
        run = p.add_run(f'{i}. {item}')
        run.font.size = Pt(12)
    else:
        run = p.add_run(f'{i}. {item}')
        run.font.size = Pt(11)

doc.add_page_break()

# ==================== 正文 ====================

# 一、执行摘要
doc.add_heading('一、执行摘要', 1)
summary = doc.add_paragraph()
summary.add_run('本策划方案旨在通过举办撕歌 APP 官方音乐人认证活动，发掘平台优质音乐人才，提升平台音乐内容质量，增强用户粘性与归属感。活动采用线上实时审核形式，由官方主持人在专属房间对报名用户进行演唱评审，通过者获得官方音乐人认证标识及相关权益。\n\n')
summary.add_run('核心亮点：').bold = True
summary.add_run('① 官方权威认证 ② 专业评审标准 ③ 丰富后续权益 ④ 持续生态建设\n')
summary.add_run('预计规模：首期 50 人参与，认证通过率 30%-50%\n')
summary.add_run('执行周期：总计 10-15 天（含准备、执行、后续跟进）')

# 二、活动背景与目的
doc.add_heading('二、活动背景与目的', 1)

doc.add_heading('2.1 活动背景', 2)
bg = doc.add_paragraph()
bg.add_run('• 平台发展需求：撕歌 APP 用户规模持续增长，需建立专业音乐人生态体系\n')
bg.add_run('• 用户需求：大量优质唱歌用户渴望获得官方认可与展示机会\n')
bg.add_run('• 行业趋势：音乐社交平台认证体系成为标配，提升平台专业度与吸引力\n')
bg.add_run('• 竞争环境：同类平台纷纷推出音乐人认证计划，需保持竞争力')

doc.add_heading('2.2 活动目的', 2)
purposes_table = doc.add_table(rows=6, cols=2)
purposes_table.style = 'Table Grid'
purposes_table.rows[0].cells[0].text = '目的维度'
purposes_table.rows[0].cells[1].text = '具体说明'
purposes_data = [
    ('人才发掘', '发掘平台内有潜力的音乐人才，建立官方音乐人库'),
    ('内容提升', '提升平台音乐内容质量，打造优质内容标杆'),
    ('用户粘性', '增强用户参与感和归属感，提高用户留存率'),
    ('品牌建设', '树立平台专业形象，提升品牌影响力'),
    ('生态构建', '构建音乐人成长体系，形成良性生态循环')
]
for i, (dim, desc) in enumerate(purposes_data, 1):
    purposes_table.rows[i].cells[0].text = dim
    purposes_table.rows[i].cells[1].text = desc

doc.add_heading('2.3 活动目标', 2)
goals = doc.add_paragraph()
goals.add_run('【量化目标】\n')
goals.add_run('• 报名人数：≥50 人（首期）\n', style='List Bullet')
goals.add_run('• 认证通过率：30%-50%（15-25 人）\n', style='List Bullet')
goals.add_run('• 活动曝光量：≥10 万次\n', style='List Bullet')
goals.add_run('• 用户满意度：≥85%\n', style='List Bullet')
goals.add_run('• 认证音乐人月活跃度：≥70%\n', style='List Bullet')

# 三、活动主题与定位
doc.add_heading('三、活动主题与定位', 1)
theme = doc.add_paragraph()
theme.add_run('活动主题：').bold = True
theme.add_run('「声动人心 · 认证启航」\n\n')
theme.add_run('活动定位：').bold = True
theme.add_run('撕歌 APP 官方权威音乐人认证计划，面向平台全体用户的常态化音乐人才选拔与认证活动\n\n')
theme.add_run('核心价值：').bold = True
theme.add_run('专业 · 公平 · 成长 · 归属')

# 四、活动概况
doc.add_heading('四、活动概况', 1)

doc.add_heading('4.1 基本信息', 2)
info_table = doc.add_table(rows=6, cols=2)
info_table.style = 'Table Grid'
info_data = [
    ('活动名称', '撕歌 APP 官方音乐人认证专场'),
    ('主办单位', '撕歌 APP 运营中心 - 音乐人生态部'),
    ('活动时间', '报名期 3 天 + 审核期 1-2 天 + 公布期 1-3 天'),
    ('活动形式', '线上实时演唱审核'),
    ('参与人数', '首期 50 人（后续可扩展）'),
    ('审核方式', '官方主持人实时评审 + 录音存档')
]
for i, (item, value) in enumerate(info_data):
    info_table.rows[i].cells[0].text = item
    info_table.rows[i].cells[1].text = value

doc.add_heading('4.2 参与对象', 2)
audience = doc.add_paragraph()
audience.add_run('• 撕歌 APP 注册用户，账号状态正常\n', style='List Bullet')
audience.add_run('• 热爱音乐，有一定演唱基础\n', style='List Bullet')
audience.add_run('• 遵守平台规则，无违规记录\n', style='List Bullet')
audience.add_run('• 愿意配合活动流程安排')

doc.add_heading('4.3 活动形式', 2)
form = doc.add_paragraph()
form.add_run('官方主持人在专属审核房间进行实时评审，参赛歌手依次进入房间演唱 1-2 分钟，主持人根据评分标准进行现场打分，通过者获得官方音乐人认证标识。审核全程录音录像存档，确保公平公正。')

# 五、目标受众分析
doc.add_heading('五、目标受众分析', 1)

audience_table = doc.add_table(rows=4, cols=3)
audience_table.style = 'Table Grid'
audience_table.rows[0].cells[0].text = '受众类型'
audience_table.rows[0].cells[1].text = '特征描述'
audience_table.rows[0].cells[2].text = '需求痛点'

audience_data = [
    ('校园歌手', '18-25 岁，学生群体，有表演欲望', '缺乏展示平台，渴望专业认可'),
    ('音乐爱好者', '20-35 岁，业余爱好，有一定基础', '希望提升技能，获得指导'),
    ('半职业歌手', '22-40 岁，有演出经验', '寻求更多发展机会，商业合作'),
    ('内容创作者', '25-35 岁，原创能力', '需要平台推广，作品曝光')
]
for i, (atype, feature, pain) in enumerate(audience_data, 1):
    audience_table.rows[i].cells[0].text = atype
    audience_table.rows[i].cells[1].text = feature
    audience_table.rows[i].cells[2].text = pain

# 六、活动流程详细规划
doc.add_heading('六、活动流程详细规划', 1)

doc.add_heading('6.1 第一阶段：前期准备（活动前 5-7 天）', 2)
prep = doc.add_paragraph()
prep.add_run('【团队组建】\n', style='List Number')
prep.add_run('• 项目负责人 1 人：整体统筹、制定时间表、协调资源、监督进度\n', style='List Bullet')
prep.add_run('• 主持人 1-2 人：熟悉审核标准、准备话术、现场把控\n', style='List Bullet')
prep.add_run('• 技术支持 1 人：房间创建、设备测试、故障处理\n', style='List Bullet')
prep.add_run('• 客服人员 2 人：群管理、答疑解惑、信息收集\n', style='List Bullet')
prep.add_run('• 宣传专员 1 人：海报设计、内容发布、数据跟踪\n', style='List Bullet')

prep.add_run('\n【资源准备】\n', style='List Number')
resources = [
    '创建官方音乐认证群（QQ 群/微信群）',
    '设计活动海报（含二维码、时间、规则）',
    '准备 FAQ 文档（常见问题解答）',
    '制作审核评分表（Excel 模板）',
    '准备备用审核房间（技术预案）',
    '配置官方认证标识（系统后台）',
    '准备录音录像设备（存档用）'
]
for res in resources:
    doc.add_paragraph(f'  • {res}', style='List Bullet')

doc.add_heading('6.2 第二阶段：宣传预热（活动前 3 天）', 2)
promo = doc.add_paragraph()
promo.add_run('Day -3：\n')
promo.add_run('  • APP 首页 Banner 上线\n', style='List Bullet')
promo.add_run('  • 官方社交媒体发布预告\n', style='List Bullet')
promo.add_run('  • 向现有音乐人群发邀请\n\n', style='List Bullet')
promo.add_run('Day -2：\n')
promo.add_run('  • 发布详细活动规则\n', style='List Bullet')
promo.add_run('  • 开放官方认证群入群通道\n', style='List Bullet')
promo.add_run('  • KOL 转发推广\n\n', style='List Bullet')
promo.add_run('Day -1：\n')
promo.add_run('  • 最后提醒报名即将开始\n', style='List Bullet')
promo.add_run('  • 群内发布报名接龙模板\n', style='List Bullet')
promo.add_run('  • 技术设备最终测试\n', style='List Bullet')

doc.add_heading('6.3 第三阶段：报名阶段（活动当天 00:00 起）', 2)
signup = doc.add_paragraph()
signup.add_run('【报名启动】\n')
signup.add_run('00:00 客服在官方认证群发布公告：\n')
signup.add_run('「【撕歌官方音乐人认证活动报名正式开启！】\n', style='Intense Quote')
signup.add_run('📅 报名时间：即日起至 [具体日期] [具体时间]\n')
signup.add_run('👥 名额限制：前 50 名（先到先得）\n')
signup.add_run('🎤 演唱要求：1-2 分钟清唱或伴奏（禁止原唱）」\n\n')

signup.add_run('【报名接龙格式】\n')
signup_format = """【撕歌音乐人认证报名】
1. 用户 ID：123456789
   用户昵称：张三
   参赛歌曲：《平凡之路》
   原唱：朴树
   
2. 用户 ID：987654321
   用户昵称：李四
   参赛歌曲：《演员》
   原唱：薛之谦"""
doc.add_paragraph(signup_format)

doc.add_heading('6.4 第四阶段：审核阶段（报名截止后 1-2 天）', 2)
audit = doc.add_paragraph()
audit.add_run('【房间准备】\n')
audit.add_run('  • 官方创建专属审核房间\n', style='List Bullet')
audit.add_run('  • 房间名称：「撕歌官方音乐人认证专场」\n', style='List Bullet')
audit.add_run('  • 房间密码：仅对报名成功用户公布\n\n', style='List Bullet')

audit.add_run('【审核流程】\n')
audit_steps = [
    '主持人开场介绍活动规则（5 分钟）',
    '按报名顺序依次邀请参赛者进入房间',
    '每位参赛者演唱 1-2 分钟',
    '主持人现场点评并记录结果',
    '审核全程录音录像存档'
]
for step in audit_steps:
    doc.add_paragraph(f'  {audit_steps.index(step)+1}. {step}', style='List Number')

audit.add_run('\n【时间控制】\n')
audit.add_run('  • 总审核时长控制在 2-3 小时内\n', style='List Bullet')
audit.add_run('  • 如遇技术问题可适当延长\n', style='List Bullet')

doc.add_heading('6.5 第五阶段：结果公布（审核结束后 1-3 天）', 2)
result = doc.add_paragraph()
result.add_run('【公布渠道】撕歌官方音乐认证群\n\n')
result.add_run('【公布内容】\n')
result.add_run('  • 通过认证的用户名单\n', style='List Bullet')
result.add_run('  • 获得的官方音乐人认证标识\n', style='List Bullet')
result.add_run('  • 后续权益说明\n', style='List Bullet')
result.add_run('  • 未通过用户的鼓励与改进建议\n', style='List Bullet')

# 七、组织架构与人员分工
doc.add_heading('七、组织架构与人员分工', 1)

org_table = doc.add_table(rows=6, cols=4)
org_table.style = 'Table Grid'
org_table.rows[0].cells[0].text = '角色'
org_table.rows[0].cells[1].text = '人数'
org_table.rows[0].cells[2].text = '核心职责'
org_table.rows[0].cells[3].text = '具体任务'

org_data = [
    ('项目负责人', '1 人', '整体统筹', '制定时间表、协调资源、监督进度、应急决策'),
    ('主持人', '1-2 人', '审核执行', '熟悉审核标准、准备话术、现场把控、点评反馈'),
    ('技术支持', '1 人', '技术保障', '房间创建、设备测试、故障处理、录音存档'),
    ('客服人员', '2 人', '用户服务', '群管理、答疑解惑、信息收集、报名统计'),
    ('宣传专员', '1 人', '推广宣传', '海报设计、内容发布、数据跟踪、效果分析')
]
for i, (role, num, duty, tasks) in enumerate(org_data, 1):
    org_table.rows[i].cells[0].text = role
    org_table.rows[i].cells[1].text = num
    org_table.rows[i].cells[2].text = duty
    org_table.rows[i].cells[3].text = tasks

# 八、宣传推广方案
doc.add_heading('八、宣传推广方案', 1)

doc.add_heading('8.1 宣传渠道', 2)
channels_table = doc.add_table(rows=5, cols=3)
channels_table.style = 'Table Grid'
channels_table.rows[0].cells[0].text = '渠道类型'
channels_table.rows[0].cells[1].text = '具体渠道'
channels_table.rows[0].cells[2].text = '预期效果'

channels_data = [
    ('APP 内', '首页 Banner、消息中心推送、开屏广告', '覆盖 100% 活跃用户'),
    ('社交媒体', '官方微博、微信公众号、抖音', '曝光量 5 万+'),
    ('社群传播', '各用户群、音乐爱好者群', '精准触达目标用户'),
    ('KOL 推广', '已认证音乐人转发', '带动粉丝参与'),
    ('线下渠道', '合作高校音乐社团', '拓展校园用户')
]
for i, (ctype, channel, effect) in enumerate(channels_data, 1):
    channels_table.rows[i].cells[0].text = ctype
    channels_table.rows[i].cells[1].text = channel
    channels_table.rows[i].cells[2].text = effect

doc.add_heading('8.2 宣传物料', 2)
materials = [
    '活动海报：主视觉海报 + 倒计时海报 + 战报模板',
    '宣传视频：往期优秀认证音乐人展示（30 秒/60 秒版本）',
    'FAQ 文档：常见问题解答（图文版）',
    'H5 页面：活动专题页（报名入口 + 规则说明）'
]
for mat in materials:
    doc.add_paragraph(f'• {mat}', style='List Bullet')

doc.add_heading('8.3 宣传节奏', 2)
rhythm = doc.add_paragraph()
rhythm.add_run('• T-7 天：内部预热，KOL 沟通\n', style='List Bullet')
rhythm.add_run('• T-3 天：官宣启动，全渠道上线\n', style='List Bullet')
rhythm.add_run('• T-1 天：倒计时提醒，制造紧迫感\n', style='List Bullet')
rhythm.add_run('• T+0 天：报名开启，实时播报进度\n', style='List Bullet')
rhythm.add_run('• T+3 天：报名截止，公布审核时间\n', style='List Bullet')
rhythm.add_run('• T+5 天：审核完成，预告结果公布\n', style='List Bullet')
rhythm.add_run('• T+7 天：结果公布，优秀选手展示\n', style='List Bullet')

# 九、预算明细表
doc.add_heading('九、预算明细表', 1)

budget_table = doc.add_table(rows=7, cols=4)
budget_table.style = 'Table Grid'
budget_table.rows[0].cells[0].text = '项目类别'
budget_table.rows[0].cells[1].text = '明细'
budget_table.rows[0].cells[2].text = '数量/天数'
budget_table.rows[0].cells[3].text = '预算（元）'

budget_data = [
    ('人力成本', '主持人', '2 人×2 天', '2000'),
    ('人力成本', '技术支持', '1 人×2 天', '1000'),
    ('人力成本', '客服人员', '2 人×3 天', '1500'),
    ('物料成本', '设计费用', '海报+H5+ 视频', '3000'),
    ('运营成本', '群管理工具', '1 个月', '500'),
    ('其他', '认证标识制作', '50 个', '1000'),
    ('合计', '', '', '9000')
]
for i, (cat, detail, qty, cost) in enumerate(budget_data, 1):
    budget_table.rows[i].cells[0].text = cat
    budget_table.rows[i].cells[1].text = detail
    budget_table.rows[i].cells[2].text = qty
    budget_table.rows[i].cells[3].text = cost

# 十、风险评估与应急预案
doc.add_heading('十、风险评估与应急预案', 1)

doc.add_heading('10.1 技术风险', 2)
tech_risk_table = doc.add_table(rows=4, cols=3)
tech_risk_table.style = 'Table Grid'
tech_risk_table.rows[0].cells[0].text = '风险项'
tech_risk_table.rows[0].cells[1].text = '概率'
tech_risk_table.rows[0].cells[2].text = '应对措施'

tech_risks = [
    ('网络不稳定', '中', '准备备用房间，允许重试一次'),
    ('设备故障', '低', '提前测试设备，技术支持待命'),
    ('系统崩溃', '低', '延后审核时间，及时通知用户'),
    ('录音失败', '低', '双设备备份录音')
]
for i, (risk, prob, measure) in enumerate(tech_risks, 1):
    tech_risk_table.rows[i].cells[0].text = risk
    tech_risk_table.rows[i].cells[1].text = prob
    tech_risk_table.rows[i].cells[2].text = measure

doc.add_heading('10.2 运营风险', 2)
ops_risk_table = doc.add_table(rows=4, cols=3)
ops_risk_table.style = 'Table Grid'
ops_risk_table.rows[0].cells[0].text = '风险项'
ops_risk_table.rows[0].cells[1].text = '概率'
ops_risk_table.rows[0].cells[2].text = '应对措施'

ops_risks = [
    ('报名超员', '中', '严格按先到先得，设置候补名单'),
    ('争议处理', '中', '建立申诉机制，保证公平公正'),
    ('用户爽约', '中', '提前确认，设置候补机制'),
    ('审核标准争议', '低', '提前公示标准，提供评分明细')
]
for i, (risk, prob, measure) in enumerate(ops_risks, 1):
    ops_risk_table.rows[i].cells[0].text = risk
    ops_risk_table.rows[i].cells[1].text = prob
    ops_risk_table.rows[i].cells[2].text = measure

doc.add_heading('10.3 舆情风险', 2)
pr_risk_table = doc.add_table(rows=3, cols=3)
pr_risk_table.style = 'Table Grid'
pr_risk_table.rows[0].cells[0].text = '风险项'
pr_risk_table.rows[0].cells[1].text = '概率'
pr_risk_table.rows[0].cells[2].text = '应对措施'

pr_risks = [
    ('负面评价', '低', '及时回应，积极沟通解决'),
    ('质疑公平性', '低', '公开评分标准，提供录音存档'),
    ('恶意攻击', '低', '建立举报机制，快速处理')
]
for i, (risk, prob, measure) in enumerate(pr_risks, 1):
    pr_risk_table.rows[i].cells[0].text = risk
    pr_risk_table.rows[i].cells[1].text = prob
    pr_risk_table.rows[i].cells[2].text = measure

# 十一、效果评估与 KPI
doc.add_heading('十一、效果评估与 KPI', 1)

kpi_table = doc.add_table(rows=6, cols=3)
kpi_table.style = 'Table Grid'
kpi_table.rows[0].cells[0].text = '评估指标'
kpi_table.rows[0].cells[1].text = '目标值'
kpi_table.rows[0].cells[2].text = '评估方法'

kpi_data = [
    ('报名人数', '≥50 人', '后台统计数据'),
    ('认证通过率', '30%-50%', '通过人数/报名人数'),
    ('活动曝光量', '≥10 万次', '各渠道数据汇总'),
    ('用户满意度', '≥85%', '问卷调查'),
    ('认证音乐人月活', '≥70%', '月度活跃数据跟踪'),
    ('内容质量提升', '优质作品 +20%', '作品评分对比')
]
for i, (kpi, target, method) in enumerate(kpi_data, 1):
    kpi_table.rows[i].cells[0].text = kpi
    kpi_table.rows[i].cells[1].text = target
    kpi_table.rows[i].cells[2].text = method

# 十二、后续权益与发展规划
doc.add_heading('十二、后续权益与发展规划', 1)

doc.add_heading('官方音乐人认证标识权益', 2)
benefits1 = [
    '个人主页显示专属认证徽章',
    '作品推荐权重提升 30%',
    '优先参与官方活动',
    '专属客服通道',
    '认证证书（电子版 + 纸质版）'
]
for benefit in benefits1:
    doc.add_paragraph(f'• {benefit}', style='List Bullet')

doc.add_heading('发展机会', 2)
benefits2 = [
    '推荐至合作音乐平台（网易云、QQ 音乐等）',
    '优先获得商业合作机会',
    '参与官方音乐制作项目',
    '定期技能培训和指导',
    '年度优秀音乐人评选资格'
]
for benefit in benefits2:
    doc.add_paragraph(f'• {benefit}', style='List Bullet')

doc.add_heading('长期规划', 2)
longterm = doc.add_paragraph()
longterm.add_run('• 季度常态化举办，形成品牌活动\n', style='List Bullet')
longterm.add_run('• 建立音乐人等级体系（铜/银/金/钻石）\n', style='List Bullet')
longterm.add_run('• 打造音乐人社区，促进交流合作\n', style='List Bullet')
longterm.add_run('• 探索商业化路径（直播打赏、商演对接等）\n', style='List Bullet')

# 附录
doc.add_page_break()
doc.add_heading('附录', 1)

doc.add_heading('附录 A：审核评分表模板', 2)
score_table = doc.add_table(rows=5, cols=4)
score_table.style = 'Table Grid'
score_table.rows[0].cells[0].text = '评分项'
score_table.rows[0].cells[1].text = '权重'
score_table.rows[0].cells[2].text = '评分 (1-10)'
score_table.rows[0].cells[3].text = '备注'

score_items = [
    ('音准节奏', '40%', '', '音准准确，节奏稳定'),
    ('音色表现', '30%', '', '音色独特，有辨识度'),
    ('情感表达', '20%', '', '情感真挚，感染力强'),
    ('整体印象', '10%', '', '台风自然，表现力佳')
]
for i, (item, weight, score, note) in enumerate(score_items, 1):
    score_table.rows[i].cells[0].text = item
    score_table.rows[i].cells[1].text = weight
    score_table.rows[i].cells[2].text = score
    score_table.rows[i].cells[3].text = note

doc.add_paragraph('\n总分：_______ / 100 分    审核结果：□通过  □不通过')
doc.add_paragraph('审核人签字：_____________    日期：_____________')

doc.add_heading('附录 B：报名接龙模板', 2)
signup_template = """【撕歌音乐人认证报名】
1. 用户 ID：___________
   用户昵称：___________
   参赛歌曲：《___________》
   原唱：___________
   联系电话/微信：___________
   
2. 用户 ID：___________
   用户昵称：___________
   参赛歌曲：《___________》
   原唱：___________
   联系电话/微信：___________"""
doc.add_paragraph(signup_template)

doc.add_heading('附录 C：常见问题 FAQ', 2)
faq = [
    ('Q1：认证是否收费？', 'A：完全免费，官方认证不收取任何费用。'),
    ('Q2：未通过可以重新报名吗？', 'A：可以，欢迎下次活动继续报名。'),
    ('Q3：认证有效期多久？', 'A：认证长期有效，但需保持活跃度。'),
    ('Q4：可以演唱原创歌曲吗？', 'A：非常欢迎，原创歌曲有额外加分。'),
    ('Q5：审核结果可以申诉吗？', 'A：可以，7 天内可向官方客服提出申诉。')
]
for q, a in faq:
    p = doc.add_paragraph()
    p.add_run(q).bold = True
    p.add_run(f'\n{a}')

# 封底
doc.add_page_break()
for _ in range(10):
    doc.add_paragraph()

closing = doc.add_paragraph()
closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
closing.add_run('—— 期待与您共同打造音乐人生态 ——\n\n')
closing.add_run('撕歌 APP 运营中心 · 音乐人生态部\n')
closing.add_run('2026 年 3 月\n')
closing.add_run('联系方式：music@singeapp.com')

for run in closing.runs:
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(100, 100, 100)

# 保存文档
output_path = '/home/admin/.openclaw/workspace/撕歌 APP 官方音乐人认证活动策划书（专业版）.docx'
doc.save(output_path)
print(f"专业版策划书已保存至：{output_path}")
